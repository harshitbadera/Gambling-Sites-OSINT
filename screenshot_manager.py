import os
import time
import asyncio
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from playwright.async_api import async_playwright
import tldextract

import config
from utils import normalize_domain, is_valid_domain

# Define output directories
SCREENSHOT_DIR = os.path.join(config.OUTPUT_DIR, "screenshots")
os.makedirs(SCREENSHOT_DIR, exist_ok=True)


async def check_liveness_and_screenshot(browser, domain: str, sem: asyncio.Semaphore) -> dict:
    """
    Check if a website is working and capture a screenshot.
    Uses a semaphore to throttle the number of concurrent pages.
    """
    async with sem:
        ext = tldextract.extract(domain)
        main_domain = f"{ext.domain}.{ext.suffix}"
        
        result = {
            "domain": domain,
            "status": "Offline",
            "screenshot_path": None,
            "response_status": None,
        }
        
        # Test HTTPS first, then fallback to HTTP
        urls_to_try = [f"https://{main_domain}", f"http://{main_domain}", f"https://{domain}", f"http://{domain}"]
        
        page = await browser.new_page()
        # Set viewport size
        await page.set_viewport_size({"width": 1280, "height": 800})
        
        for url in urls_to_try:
            try:
                # 12 seconds timeout to avoid hanging the pipeline
                response = await page.goto(url, timeout=12000, wait_until="load")
                
                # Check if we got a valid response (HTTP < 400 is ideal, but even higher codes mean the server is technically up)
                if response:
                    status_code = response.status
                    result["response_status"] = status_code
                    
                    if status_code < 400 or status_code in [401, 403]: # 401/403 means site is working but restricted
                        # Wait 2 seconds for JS/assets to render
                        await asyncio.sleep(2)
                        
                        # Save screenshot
                        ss_filename = f"{domain}_{int(time.time())}.png"
                        ss_path = os.path.join(SCREENSHOT_DIR, ss_filename)
                        await page.screenshot(path=ss_path)
                        
                        result["status"] = "Working"
                        result["screenshot_path"] = ss_path
                        break
            except Exception:
                continue
        
        await page.close()
        return result


async def run_screenshot_pipeline(domains: list) -> list:
    """
    Execute Playwright concurrent screenshot jobs.
    """
    # Limit to 5 concurrent browser tabs to avoid memory/CPU spikes
    sem = asyncio.Semaphore(5)
    
    async with async_playwright() as p:
        # Launch headless browser
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-web-security", "--ignore-certificate-errors"]
        )
        
        tasks = [check_liveness_and_screenshot(browser, d, sem) for d in domains]
        results = await asyncio.gather(*tasks)
        
        await browser.close()
        return results


def generate_word_report(results: list, output_filename: str = None) -> str:
    """
    Generate a formatted Word document (.docx) containing the screenshots of working domains
    and the status of non-working domains.
    """
    if not output_filename:
        output_filename = f"liveness_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
    doc_path = os.path.join(config.OUTPUT_DIR, "temp", output_filename)
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    
    doc = Document()
    
    # Page styling: Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Section
    title = doc.add_paragraph()
    title_run = title.add_run("GAMBLING OSINT — SITE LIVENESS REPORT")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Report Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                            f"Total domains checked: {len(results)}\n"
                            f"Working sites: {len([r for r in results if r['status'] == 'Working'])}\n"
                            f"Offline sites: {len([r for r in results if r['status'] == 'Offline'])}")
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(10)
    meta_run.italic = True
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(20)
    
    # Add a summary table
    doc.add_heading("1. Summary of Checked Domains", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Domain'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'HTTP Status Code'
    
    # Formatting header row font
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Arial"
                
    for r in sorted(results, key=lambda x: x["domain"]):
        row_cells = table.add_row().cells
        row_cells[0].text = r["domain"]
        row_cells[1].text = r["status"]
        row_cells[2].text = str(r["response_status"]) if r["response_status"] else "-"
        
        # Color coding cell backgrounds based on status
        if r["status"] == "Working":
            # Light green background
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E2F0D9"/>')
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
        else:
            # Light red background
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FCE4D6"/>')
            row_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            
    doc.add_paragraph("").paragraph_format.space_after = Pt(20)
    
    # Detailed Section with screenshots
    doc.add_heading("2. Site Screenshots & Details", level=1)
    
    for r in sorted(results, key=lambda x: x["domain"]):
        # Add Page Break for each domain to look neat
        doc.add_paragraph("").paragraph_format.space_after = Pt(10)
        
        h = doc.add_heading(f"Site: {r['domain']}", level=2)
        h.paragraph_format.space_before = Pt(15)
        
        status_para = doc.add_paragraph()
        s_label = status_para.add_run("Status: ")
        s_label.font.bold = True
        
        s_val = status_para.add_run(r["status"])
        s_val.font.bold = True
        
        if r["status"] == "Working":
            s_val.font.color.rgb = RGBColor(46, 117, 89) # Dark Green
            
            http_para = doc.add_paragraph()
            h_label = http_para.add_run("HTTP Status Code: ")
            h_label.font.bold = True
            http_para.add_run(str(r["response_status"]))
            
            # Insert screenshot
            if r["screenshot_path"] and os.path.exists(r["screenshot_path"]):
                p = doc.add_paragraph()
                p.alignment = 1 # Center align
                run = p.add_run()
                run.add_picture(r["screenshot_path"], width=Inches(5.8))
                
                cap = doc.add_paragraph()
                cap.alignment = 1
                c_run = cap.add_run(f"Figure: Screenshot of {r['domain']}")
                c_run.font.size = Pt(9)
                c_run.italic = True
        else:
            s_val.font.color.rgb = RGBColor(192, 0, 0) # Dark Red
            offline_para = doc.add_paragraph()
            offline_run = offline_para.add_run("[OFFLINE] The website did not respond or failed to load. No screenshot is available.")
            offline_run.italic = True
            
        doc.add_paragraph("").paragraph_format.space_after = Pt(20)
        
    doc.save(doc_path)
    return doc_path


def run_liveness_and_generate_report(domains: list, output_filename: str = None) -> str:
    """
    Synchronous wrapper to run the async Playwright pipeline
    and generate the Word document.
    """
    # Create or get event loop
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
    results = loop.run_until_complete(run_screenshot_pipeline(domains))
    report_path = generate_word_report(results, output_filename)
    return report_path
